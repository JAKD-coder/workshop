"""
Generate Reflex Connect Monthly Service Report as a Word document and export to PDF.
Usage: python3 generate_report.py --month "April 2026"
"""

import argparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess
import os
from datetime import datetime

TEMPLATE_PATH = "templates/monthly_service_report_template.docx"


def build_template():
    """Create and save the Word template."""
    doc = Document()

    # Page margins
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

    def add_heading(text, level=1, color=RGBColor(0x1F, 0x49, 0x7D)):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = color
        return p

    def add_placeholder(label):
        p = doc.add_paragraph()
        run = p.add_run(f"[{label}]")
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.font.italic = True
        return p

    # Cover
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Reflex Connect")
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("Monthly Service Report")
    run2.font.size = Pt(18)

    month_p = doc.add_paragraph()
    month_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_placeholder("Month Year")

    doc.add_page_break()

    # 1. Executive Summary
    add_heading("1. Executive Summary")
    add_placeholder("Brief overview of service performance for the month")

    # 2. Service Performance
    add_heading("2. Service Performance")

    add_heading("2.1 Uptime & Availability", level=2)
    table = doc.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Service"
    hdr[1].text = "Target SLA"
    hdr[2].text = "Actual"
    row = table.rows[1].cells
    row[0].text = "[Service Name]"
    row[1].text = "99.9%"
    row[2].text = "[x.xx%]"
    doc.add_paragraph()

    add_heading("2.2 Incident Summary", level=2)
    inc_table = doc.add_table(rows=2, cols=4)
    inc_table.style = "Table Grid"
    ih = inc_table.rows[0].cells
    ih[0].text = "Incident ID"
    ih[1].text = "Date"
    ih[2].text = "Description"
    ih[3].text = "Resolution"
    ir = inc_table.rows[1].cells
    ir[0].text = "[ID]"
    ir[1].text = "[Date]"
    ir[2].text = "[Description]"
    ir[3].text = "[Resolution]"
    doc.add_paragraph()

    # 3. Contact Centre Metrics
    add_heading("3. Contact Centre Metrics")
    cc_table = doc.add_table(rows=2, cols=4)
    cc_table.style = "Table Grid"
    ch = cc_table.rows[0].cells
    ch[0].text = "Metric"
    ch[1].text = "Target"
    ch[2].text = "Actual"
    ch[3].text = "Trend"
    cr = cc_table.rows[1].cells
    cr[0].text = "[Metric Name]"
    cr[1].text = "[Target]"
    cr[2].text = "[Actual]"
    cr[3].text = "[↑/↓/→]"
    doc.add_paragraph()

    # 4. Platform & Infrastructure
    add_heading("4. Platform & Infrastructure")
    add_placeholder("Summary of platform changes, upgrades, or issues")

    # 5. Client Summary (Capitec, FNB, Vodacom, etc.)
    add_heading("5. Client Summary")
    for client in ["Capitec", "FNB", "Vodacom"]:
        add_heading(f"5.x {client}", level=2)
        add_placeholder(f"Key activities and issues for {client} this month")

    # 6. Financial Summary
    add_heading("6. Financial Summary")
    add_placeholder("Revenue, costs, and budget variance for the month")

    # 7. Actions & Next Steps
    add_heading("7. Actions & Next Steps")
    action_table = doc.add_table(rows=2, cols=3)
    action_table.style = "Table Grid"
    ah = action_table.rows[0].cells
    ah[0].text = "Action"
    ah[1].text = "Owner"
    ah[2].text = "Due Date"
    ar = action_table.rows[1].cells
    ar[0].text = "[Action Item]"
    ar[1].text = "[Owner]"
    ar[2].text = "[Date]"
    doc.add_paragraph()

    # 8. Appendix
    add_heading("8. Appendix")
    add_placeholder("Supporting data, charts, or raw metrics")

    os.makedirs("templates", exist_ok=True)
    doc.save(TEMPLATE_PATH)
    print(f"Template saved: {TEMPLATE_PATH}")
    return doc


def generate_report(month: str):
    """Generate a filled report from the template and export to PDF."""
    if not os.path.exists(TEMPLATE_PATH):
        print("Template not found, building it first...")
        build_template()

    doc = Document(TEMPLATE_PATH)

    # Replace month placeholder
    for para in doc.paragraphs:
        if "[Month Year]" in para.text:
            for run in para.runs:
                run.text = run.text.replace("[Month Year]", month)

    output_docx = f"output/Reflex_Connect_Service_Report_{month.replace(' ', '_')}.docx"
    os.makedirs("output", exist_ok=True)
    doc.save(output_docx)
    print(f"Report saved: {output_docx}")

    # Export to PDF via LibreOffice (if available)
    result = subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", "output", output_docx],
        capture_output=True, text=True
    )
    if result.returncode == 0:
        pdf_path = output_docx.replace(".docx", ".pdf")
        print(f"PDF exported: {pdf_path}")
    else:
        print("PDF export requires LibreOffice. Install with: brew install --cask libreoffice")

    return output_docx


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Generate Reflex Connect Monthly Service Report")
    parser.add_argument("--month", default=datetime.now().strftime("%B %Y"), help="Report month (e.g. 'April 2026')")
    parser.add_argument("--build-template", action="store_true", help="Rebuild the Word template only")
    args = parser.parse_args()

    if args.build_template:
        build_template()
    else:
        generate_report(args.month)
